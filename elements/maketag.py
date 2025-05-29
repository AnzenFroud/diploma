from docx import Document
import tkinter as tk
from tkinter import messagebox
import datetime
import os
import platform
import subprocess
from tkinter import ttk


def generate_sticker(table):
    row_index = table.getSelectedRow()
    if row_index is None:
        messagebox.showwarning("Выбор строки", "Пожалуйста, выделите строку в таблице.")
        return

    row_data = table.model.getRecordAtRow(row_index)
    all_columns = list(row_data.keys())

    selector_win = tk.Toplevel()
    selector_win.title("Создание этикетки")
    selector_win.geometry("600x500")

    # Верхняя рамка для селекторов
    selector_frame = tk.LabelFrame(selector_win, text="Выберите столбцы для этикетки", padx=10, pady=10)
    selector_frame.pack(fill="x", padx=10, pady=5)

    column_selectors = []

    def add_selector():
        var = tk.StringVar()
        var.set(all_columns[0])

        row = tk.Frame(selector_frame)
        row.pack(fill="x", pady=2)

        label = tk.Label(row, text=f"Поле {len(column_selectors)+1}:", width=15, anchor="w")
        label.pack(side="left")

        dropdown = ttk.Combobox(row, textvariable=var, values=all_columns, state="readonly", width=40)
        dropdown.pack(side="left")

        column_selectors.append(var)
        update_preview()

    # Предпросмотр этикетки
    preview_frame = tk.LabelFrame(selector_win, text="Предпросмотр этикетки", padx=10, pady=10)
    preview_frame.pack(fill="both", expand=True, padx=10, pady=5)

    preview_text = tk.Text(preview_frame, height=10, wrap="word", state="disabled")
    preview_text.pack(fill="both", expand=True)

    def update_preview():
        selected = [v.get() for v in column_selectors]
        preview_text.configure(state="normal")
        preview_text.delete("1.0", tk.END)

        if selected:
            preview_text.insert(tk.END, "Этикетка оборудования\n\n")
            for col in selected:
                preview_text.insert(tk.END, f"{col}: {row_data[col]}\n")
        else:
            preview_text.insert(tk.END, "Нет выбранных полей")

        preview_text.configure(state="disabled")

    def create_label():
        selected = [v.get() for v in column_selectors]
        if not selected:
            messagebox.showwarning("Ошибка", "Выберите хотя бы один столбец.")
            return

        doc = Document()
        doc.add_heading("Этикетка оборудования", level=1)

        for col in selected:
            doc.add_paragraph(f"{col}: {row_data[col]}")

        today = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        filename = f"etiketka_{today}.docx"
        filepath = os.path.join(os.getcwd(), filename)
        doc.save(filepath)

        try:
            if platform.system() == "Darwin":
                subprocess.call(["open", filepath])
            elif platform.system() == "Windows":
                os.startfile(filepath)
            elif platform.system() == "Linux":
                subprocess.call(["xdg-open", filepath])
        except Exception as e:
            messagebox.showinfo("Готово", f"Этикетка создана, но не удалось открыть файл:\n{filepath}\n\n{e}")
            selector_win.destroy()
            return

        messagebox.showinfo("Готово", f"Этикетка успешно создана:\n{filepath}")
        selector_win.destroy()

    # Кнопки управления
    button_frame = tk.Frame(selector_win)
    button_frame.pack(pady=10)

    add_btn = tk.Button(button_frame, text="Добавить столбец", command=add_selector)
    add_btn.pack(side="left", padx=5)

    generate_btn = tk.Button(button_frame, text="Создать этикетку", command=create_label)
    generate_btn.pack(side="left", padx=5)

    # Добавить хотя бы один по умолчанию
    add_selector()
