from docx import Document
import tkinter as tk
from tkinter import messagebox
import datetime
import os
import platform
import subprocess

def generate_sticker(table):
    row_index = table.getSelectedRow()
    if row_index is None:
        messagebox.showwarning("Выбор строки", "Пожалуйста, выделите строку в таблице.")
        return

    row_data = table.model.getRecordAtRow(row_index)
    all_columns = list(row_data.keys())

    # Окно выбора столбцов
    selector_win = tk.Toplevel()
    selector_win.title("Выберите столбцы для этикетки")
    selector_win.geometry("400x300")

    column_selectors = []

    def add_selector():
        var = tk.StringVar()
        var.set(all_columns[0])
        dropdown = tk.OptionMenu(selector_win, var, *all_columns)
        dropdown.pack(pady=2)
        column_selectors.append(var)

    def create_label():
        selected = [v.get() for v in column_selectors]
        if not selected:
            messagebox.showwarning("Ошибка", "Выберите хотя бы один столбец.")
            return

        doc = Document()
        doc.add_heading("Этикетка оборудования", level=1)

        for col in selected:
            doc.add_paragraph(f"{col}: {row_data[col]}")

        today = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        filename = f"etiketka_{today}.docx"
        filepath = os.path.join(os.getcwd(), filename)
        doc.save(filepath)

        try:
            if platform.system() == "Darwin":
                subprocess.call(["open", filepath])
            elif platform.system() == "Windows":
                os.startfile(filepath)
            elif platform.system() == "Linux":
                subprocess.call(["xdg-open", filepath])
        except Exception as e:
            messagebox.showinfo("Готово", f"Этикетка создана, но не удалось открыть файл:\n{filepath}\n\n{e}")
            selector_win.destroy()
            return

        messagebox.showinfo("Готово", f"Этикетка успешно создана:\n{filepath}")
        selector_win.destroy()

    # Кнопка для добавления селекторов
    add_btn = tk.Button(selector_win, text="Добавить столбец", command=add_selector)
    add_btn.pack(pady=5)

    # Кнопка создания этикетки
    generate_btn = tk.Button(selector_win, text="Создать этикетку", command=create_label)
    generate_btn.pack(pady=10)

    # Добавить хотя бы один по умолчанию
    add_selector()
