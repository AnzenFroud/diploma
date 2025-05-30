from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import datetime
import os
import platform
import subprocess

def generate_act(table):
    row_index = table.getSelectedRow()
    if row_index is None:
        messagebox.showwarning("Выбор строки", "Пожалуйста, выделите строку в таблице.")
        return

    row_data = table.model.getRecordAtRow(row_index)

    act_window = tk.Toplevel()
    act_window.title("Создание акта")
    act_window.geometry("500x300")
    act_window.resizable(False, False)

    act_types = ['Акт приема-передачи', 'Акт осмотра', 'Акт списания']
    act_type_var = tk.StringVar(value=act_types[0])
    path_var = tk.StringVar()

    def select_path():
        path = filedialog.askdirectory(initialdir=os.getcwd(), title="Выберите папку для сохранения")
        if path:
            path_var.set(path)

    # Тип акта
    ttk.Label(act_window, text="Тип акта:", font=('Segoe UI', 10)).pack(pady=(20, 2))
    ttk.OptionMenu(act_window, act_type_var, act_types[0], *act_types).pack()

    # Путь
    ttk.Label(act_window, text="Путь для сохранения:", font=('Segoe UI', 10)).pack(pady=(15, 2))
    ttk.Entry(act_window, textvariable=path_var, width=50).pack(pady=(0, 5))
    ttk.Button(act_window, text="Выбрать папку", command=select_path).pack()

    # Кнопка создания
    def create_act():
        save_path = path_var.get()
        if not save_path:
            messagebox.showwarning("Ошибка", "Пожалуйста, выберите путь для сохранения.")
            return

        doc = Document()
        today = datetime.datetime.now().strftime("%Y-%m-%d")
        act_type = act_type_var.get()

        heading = doc.add_heading(act_type, level=1)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        date_paragraph = doc.add_paragraph(f"Дата: {today}")
        date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_paragraph()

        if act_type == 'Акт приема-передачи':
            doc.add_paragraph("Настоящий акт составлен о том, что:").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.add_paragraph()
            for col, val in row_data.items():
                para = doc.add_paragraph(f"{col}: {val}")
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                para.runs[0].font.size = Pt(11)
            doc.add_paragraph()
            doc.add_paragraph("Оборудование было передано от __________ к __________.").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.add_paragraph()
            doc.add_paragraph("Подписи сторон:\nПолучатель: __________     Передающая сторона: __________").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            doc.add_paragraph("Данные о строке:").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for col, val in row_data.items():
                para = doc.add_paragraph(f"{col}: {val}")
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                para.runs[0].font.size = Pt(11)

        filename = f"{act_type.replace(' ', '_').lower()}_{today}.docx"
        filepath = os.path.join(save_path, filename)
        doc.save(filepath)

        try:
            if platform.system() == "Darwin":
                subprocess.call(["open", filepath])
            elif platform.system() == "Windows":
                os.startfile(filepath)
            elif platform.system() == "Linux":
                subprocess.call(["xdg-open", filepath])
        except Exception as e:
            messagebox.showinfo("Готово", f"Акт создан, но не удалось открыть файл:\n{filepath}\n\n{e}")
            act_window.destroy()
            return

        messagebox.showinfo("Готово", f"Акт успешно создан:\n{filepath}")
        act_window.destroy()

    ttk.Button(act_window, text="Создать акт", command=create_act).pack(pady=20)
